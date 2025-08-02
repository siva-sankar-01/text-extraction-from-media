from flask import Flask, request, jsonify
from flask_cors import CORS
import cv2
import os
import math
import keras_ocr
import docx

app = Flask(__name__)
CORS(app)

@app.route('/video_convert', methods=['POST'])
def convert_video():
    try:
        video_src = request.files['video']
        if not os.path.exists('videos'):
            os.makedirs('videos')
        video_path = os.path.join('videos', video_src.filename)
        video_src.save(video_path)

        cap = cv2.VideoCapture(video_path)
        if not os.path.exists('images'):
            os.makedirs('images')

        for file in os.listdir('images'):
            os.remove(os.path.join('images', file))

        currentframe = 0
        while True:
            ret, frame = cap.read()
            if not ret:
                break
            frame_path = f'images/frame{currentframe}.jpg'
            cv2.imwrite(frame_path, frame)
            currentframe += 1

        pipeline = keras_ocr.pipeline.Pipeline()
        output_text = ""

        for i in range(currentframe):
            image_path = f'images/frame{i}.jpg'
            if os.path.exists(image_path):
                image = keras_ocr.tools.read(image_path)
                predictions = pipeline.recognize([image])
                for prediction in predictions[0]:
                    output_text += prediction[0] + " "

        doc = docx.Document()
        doc.add_paragraph(output_text)
        doc.save("Batch14.docx")

        return jsonify({'data': output_text})

    except Exception as e:
        return jsonify({'error': str(e)})

@app.route('/image_convert', methods=['POST'])
def image_video():
    try:
        image_src = request.files['image']
        if not os.path.exists('images1'):
            os.makedirs('images1')
        image_path = os.path.join('images1', image_src.filename)
        image_src.save(image_path)

        pipeline = keras_ocr.pipeline.Pipeline()
        image = keras_ocr.tools.read(image_path)
        predictions = pipeline.recognize([image])

        output_text = ""
        for prediction in predictions[0]:
            output_text += prediction[0] + " "

        doc = docx.Document()
        doc.add_paragraph(output_text)
        doc.save("Batch1.docx")

        return jsonify({'data': output_text})

    except Exception as e:
        return jsonify({'error': str(e)})

if __name__ == '__main__':
    app.run(debug=True)
